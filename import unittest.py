import unittest
from unittest.mock import patch, mock_open, MagicMock
import os
from docx import Document

# Import fungsi yang akan diuji
from DemoTugas import tambah_pesanan, tampilkan_pesanan, update_pesanan, hapus_pesanan, cari_pesanan, DATA_FILE

class TestRestaurantManagement(unittest.TestCase):

    def setUp(self):
        # Membuat dokumen uji
        self.test_file = "test_pesanan.docx"
        doc = Document()
        doc.add_heading("Data Pesanan Restoran", level=1)
        doc.save(self.test_file)

        # Ganti nama file ke dokumen uji
        global DATA_FILE
        DATA_FILE = self.test_file

    def tearDown(self):
        # Hapus dokumen uji setelah selesai
        if os.path.exists(self.test_file):
            os.remove(self.test_file)

    @patch("builtins.input", side_effect=["John", "Pizza", "2", "path/to/image.jpg"])
    @patch("os.path.exists", return_value=True)
    @patch("restaurant_management.Document")
    def test_tambah_pesanan(self, mock_document, mock_exists, mock_input):
        mock_doc = MagicMock()
        mock_document.return_value = mock_doc

        tambah_pesanan()

        # Memastikan add_paragraph dan add_picture dipanggil
        mock_doc.add_paragraph.assert_called_with("Nama: John, Menu: Pizza, Jumlah: 2, Status: Diproses")
        mock_doc.add_picture.assert_called_with("path/to/image.jpg", width=mock_doc.add_picture.call_args[1]['width'])
        mock_doc.save.assert_called_with(self.test_file)

    @patch("restaurant_management.Document")
    def test_tampilkan_pesanan(self, mock_document):
        mock_doc = MagicMock()
        mock_doc.paragraphs = [
            MagicMock(text="Data Pesanan Restoran"),
            MagicMock(text="Nama: John, Menu: Pizza, Jumlah: 2, Status: Diproses"),
        ]
        mock_document.return_value = mock_doc

        with patch("builtins.print") as mock_print:
            tampilkan_pesanan()

            # Memastikan output ditampilkan
            mock_print.assert_any_call("\nData Pesanan:")
            mock_print.assert_any_call("Nama: John, Menu: Pizza, Jumlah: 2, Status: Diproses")

    @patch("builtins.input", side_effect=["John", "Selesai"])
    @patch("restaurant_management.Document")
    def test_update_pesanan(self, mock_document, mock_input):
        mock_doc = MagicMock()
        mock_doc.paragraphs = [
            MagicMock(text="Nama: John, Menu: Pizza, Jumlah: 2, Status: Diproses"),
        ]
        mock_document.return_value = mock_doc

        update_pesanan()

        # Memastikan teks diperbarui
        self.assertIn("Selesai", mock_doc.paragraphs[0].text)
        mock_doc.save.assert_called_with(self.test_file)

    @patch("builtins.input", side_effect=["John"])
    @patch("restaurant_management.Document")
    def test_hapus_pesanan(self, mock_document, mock_input):
        mock_doc = MagicMock()
        mock_doc.paragraphs = [
            MagicMock(text="Nama: John, Menu: Pizza, Jumlah: 2, Status: Batal"),
            MagicMock(text="Nama: Jane, Menu: Burger, Jumlah: 1, Status: Diproses"),
        ]
        mock_document.return_value = mock_doc

        hapus_pesanan()

        # Memastikan paragraf yang mengandung 'Batal' dihapus
        self.assertNotIn("Nama: John, Menu: Pizza, Jumlah: 2, Status: Batal", [p.text for p in mock_doc.paragraphs])
        mock_doc.save.assert_called_with(self.test_file)

    @patch("builtins.input", side_effect=["Jane"])
    @patch("restaurant_management.Document")
    def test_cari_pesanan(self, mock_document, mock_input):
        mock_doc = MagicMock()
        mock_doc.paragraphs = [
            MagicMock(text="Nama: John, Menu: Pizza, Jumlah: 2, Status: Diproses"),
            MagicMock(text="Nama: Jane, Menu: Burger, Jumlah: 1, Status: Diproses"),
        ]
        mock_document.return_value = mock_doc

        with patch("builtins.print") as mock_print:
            cari_pesanan()

            # Memastikan output yang benar ditampilkan
            mock_print.assert_any_call("Ditemukan: Nama: Jane, Menu: Burger, Jumlah: 1, Status: Diproses")

if __name__ == "__main__":
    unittest.main()
